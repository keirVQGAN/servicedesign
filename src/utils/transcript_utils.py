from docx import Document
from docx.shared import Pt
from pathlib import Path
import re
import string
import nltk
import json
import os

nltk.download('punkt')
nltk.download('stopwords')

from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize


def save_tutorial(response, student):

    folder_name = f'./output/tutorials/{student}'
    file_name = f'{folder_name}/{student}.json'
    # Load existing data or create an empty list
    try:
        with open(file_name, 'r') as file:
            existing_data = json.load(file)
            if not isinstance(existing_data, list):
                existing_data = []
    except (FileNotFoundError, json.JSONDecodeError):
        existing_data = []

    # Append response and save to json file
    existing_data.append(response)
    with open(file_name, 'w') as file:
        json.dump(existing_data, file, indent=4)

    # Save response text to txt file
    response_text = response['response']
    with open(f'{folder_name}/{student}.txt', 'w') as file:
        file.write(response_text)

    # Append response and student data to master_tutorials.json
    master_file = './output/tutorials/master_tutorials.json'

    try:
        with open(master_file, 'r') as file:
            master_data = json.load(file)
            if not isinstance(master_data, list):
                master_data = []
    except (FileNotFoundError, json.JSONDecodeError):
        master_data = []

    master_data.append({'student': student, 'response': response})
    with open(master_file, 'w') as file:
        json.dump(master_data, file, indent=4)

        
def clean_text(input_file, output_file, student):
    with open(input_file, 'r') as file:
        content = file.readlines()

    # This regex pattern should match most timecode formats
    timecode_pattern = r"\d+:\d+:\d+.\d+ --> \d+:\d+:\d+.\d+"

    cleaned_lines = [line for line in content if not re.match(timecode_pattern, line.strip())]

    cleaned_content = ' '.join(cleaned_lines)

    # Replace names
    cleaned_content = cleaned_content.replace('Keir Williams', 'Keir')
    cleaned_content = cleaned_content.replace(student, 'Student')

    # tokenize the text
    tokens = word_tokenize(cleaned_content)

    # remove punctuation
    table = str.maketrans('', '', string.punctuation)
    stripped = [w.translate(table) for w in tokens]

    # remove stop words
    stop_words = set(stopwords.words('english'))
    words = [word for word in stripped if not word in stop_words]

    cleaned_text = ' '.join(words)

    with open(output_file, 'w') as file:
        file.write(cleaned_text)

    return cleaned_text


def txt2docx(input_path):
    for txt_file in Path(input_path).rglob('*.txt'):
        if 'clean' in txt_file.name: continue

        docx_file = txt_file.with_suffix('.docx')
        document = Document()

        for i, line in enumerate(txt_file.read_text().strip().split('\n')):
            line = line.strip()
            if not line: continue

            p = document.add_paragraph()
            run = p.add_run(line[:-1] if line.endswith(':') else line)

            run.bold = line.endswith(':') or i == 0  # set bold if it's first line or ends with ':'
            run.font.name = 'Arial'
            run.font.size = Pt(12) if run.bold else Pt(10)
            p.paragraph_format.space_after = Pt(6)

        document.save(docx_file)
